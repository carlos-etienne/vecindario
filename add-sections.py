"""
Agrega las secciones 15 (Roadmap dashboard) y 16 (Estrategia producto)
al brief-diseno-vecindario.docx SIN regenerar el documento.
Preserva todas las ediciones existentes de Carlos.

Estrategia:
- Abrir el .docx actual
- Construir los nuevos elementos (párrafos + tablas) con python-docx
- Insertarlos antes del párrafo de cierre ("Tómate tu tiempo...")
- Guardar
"""

from docx import Document
from docx.shared import Pt, RGBColor, Twips, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy

DOC_PATH = "brief-diseno-vecindario.docx"

# Paleta (debe coincidir con generate-brief.js)
P = {
    "primary":   "2A1F17",
    "body":      "2B2520",
    "secondary": "6B5D52",
    "accent":    "B08050",
    "surface":   "FAF6F0",
    "divider":   "D6C7B2",
    "need":      "8B5A2B",
    "have":      "5A6B4A",
    "placeholder": "8A7A6A",
}
FONT = "Calibri"


def set_run(run, text, size=22, color=P["body"], bold=False, italic=False):
    run.text = text
    run.font.name = FONT
    run.font.size = Pt(size / 2)  # size is in half-points in docx-js; python-docx uses pt
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    # set eastAsia font too
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)


def _apply_heading_style(p, style_id):
    """Set pStyle by raw styleId (works around python-docx locale issues)."""
    pPr = p._element.get_or_add_pPr()
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), style_id)


# Style IDs are localized to Spanish in this docx: Ttulo1, Ttulo2, Ttulo3
H1_ID = "Ttulo1"
H2_ID = "Ttulo2"
H3_ID = "Ttulo3"


def add_h1(doc, text, page_break_before=False):
    """Heading 1 with accent bottom border"""
    p = doc.add_paragraph()
    _apply_heading_style(p, H1_ID)
    if page_break_before:
        pPr = p._element.get_or_add_pPr()
        pbd = OxmlElement("w:pageBreakBefore")
        pPr.append(pbd)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(P["primary"])
    # accent bottom border
    pPr = p._element.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), P["accent"])
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def add_h2(doc, text, tag=None):
    """Heading 2, optionally with a tag pill appended"""
    p = doc.add_paragraph()
    _apply_heading_style(p, H2_ID)
    run = p.add_run(text + "   ")
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(P["primary"])
    if tag:
        tag_text = "  ✅  YA LO SÉ  " if tag == "have" else "  ✏️  NECESITO DE TÍ  "
        tag_color = P["have"] if tag == "have" else P["need"]
        trun = p.add_run(tag_text)
        trun.bold = True
        trun.font.name = FONT
        trun.font.size = Pt(8)
        trun.font.color.rgb = RGBColor.from_string(tag_color)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    _apply_heading_style(p, H3_ID)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(P["accent"])
    return p


def add_body(doc, parts):
    """parts: list of {"text":..., "bold":..., "italic":..., "color":...} or plain string"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.3
    pf.space_after = Pt(7)
    if isinstance(parts, str):
        parts = [parts]
    for part in parts:
        if isinstance(part, str):
            run = p.add_run(part)
            set_run(run, part)
        elif isinstance(part, dict):
            text = part["text"]
            run = p.add_run(text)
            set_run(run, text,
                    color=part.get("color", P["body"]),
                    bold=part.get("bold", False),
                    italic=part.get("italic", False))
        else:
            text, opts = part
            run = p.add_run(text)
            set_run(run, text,
                    color=opts.get("color", P["body"]),
                    bold=opts.get("bold", False),
                    italic=opts.get("italic", False))
    return p


def add_bullet(doc, parts, level=0):
    """Bullet using existing numId=2 (defined in the document)."""
    p = doc.add_paragraph()
    # Apply numbering: numId=2, ilvl=level
    pPr = p._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "2")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    # indentation matching the rest of the doc
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(720 + level * 720))
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)

    pf = p.paragraph_format
    pf.line_spacing = 1.25
    pf.space_after = Pt(4)
    if isinstance(parts, str):
        parts = [parts]
    for part in parts:
        if isinstance(part, str):
            run = p.add_run(part)
            set_run(run, part)
        elif isinstance(part, dict):
            text = part["text"]
            run = p.add_run(text)
            set_run(run, text,
                    color=part.get("color", P["body"]),
                    bold=part.get("bold", False),
                    italic=part.get("italic", False))
        else:
            text, opts = part
            run = p.add_run(text)
            set_run(run, text,
                    color=opts.get("color", P["body"]),
                    bold=opts.get("bold", False),
                    italic=opts.get("italic", False))
    return p


def add_placeholder(doc, text, lines=1):
    """Dotted-underlined placeholder line(s)"""
    out = []
    for i in range(lines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_after = Pt(3)
        # dotted bottom border
        pPr = p._element.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "dotted")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "8")
        bottom.set(qn("w:color"), P["divider"])
        pbdr.append(bottom)
        pPr.append(pbdr)
        run_text = text if i == 0 else "\u00A0"
        run = p.add_run(run_text)
        set_run(run, run_text, color=P["placeholder"], italic=True)
        out.append(p)
    return out


def add_gap(doc, twips=120):
    """Empty paragraph with spacing"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(twips / 20)
    return p


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    """margins in twips (1/20 pt)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def add_data_table(doc, headers, rows, col_widths=None):
    """Styled table matching the brief's design"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, P["primary"])
        set_cell_margins(cell)
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
    # set header row to repeat
    trPr = table.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            set_cell_margins(cell)
            if ri % 2 == 0:
                shade_cell(cell, P["surface"])
            cell.paragraphs[0].text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.name = FONT
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string(P["body"] if ci > 0 else P["primary"])
            if ci == 0:
                run.bold = True
        # cantSplit on data rows
        trPr = table.rows[ri + 1]._tr.get_or_add_trPr()
        cantSplit = OxmlElement("w:cantSplit")
        trPr.append(cantSplit)
    # column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    # borders
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "bottom"]:
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), P["accent"])
        tblBorders.append(b)
    for border_name in ["left", "right"]:
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    ih = OxmlElement("w:insideH")
    ih.set(qn("w:val"), "single")
    ih.set(qn("w:sz"), "2")
    ih.set(qn("w:color"), P["divider"])
    tblBorders.append(ih)
    iv = OxmlElement("w:insideV")
    iv.set(qn("w:val"), "nil")
    tblBorders.append(iv)
    tblPr.append(tblBorders)
    return table


def build_section_15(doc):
    """Sección 15: Roadmap futuro - dashboard operacional"""
    add_h1(doc, "15 · Roadmap futuro: dashboard operacional", page_break_before=True)
    add_body(doc, [
        {"text": "Esto no es para el MVP, pero lo documento aquí para que el backend se diseñe desde el inicio ", "bold": True, "color": P["accent"]},
        {"text": "con esta capacidad en mente", "bold": True},
        " — y no haya que rediseñar después.",
    ])
    add_body(doc, [
        "La idea: extender el panel admin hasta convertirlo en un ",
        {"text": "centro de operaciones del bar", "bold": True},
        ". Las funciones siguen el patrón ",
        {"text": "capturar dato (desde donde sea) → guardarlo en D1 → verlo en dashboard / notificar", "italic": True},
        ". Cada función nueva es agregar un comando más al mismo sistema, no construir algo aparte.",
    ])

    # Perfiles de usuario
    add_h2(doc, "Perfiles de usuario previstos")
    add_body(doc, "El dashboard pasa de 1 solo usuario a varios perfiles. Esto fuerza que la autenticación se diseñe multi-usuario con roles desde el inicio (aunque al principio solo exista tu cuenta).")
    add_data_table(doc,
        ["Perfil", "Qué puede hacer", "Cómo accede"],
        [
            ["Tú", "Todo (admin, configuración, todos los módulos)", "Dashboard + Telegram"],
            ["Socio", "Gastos, ver reportes (vistas limitadas)", "Dashboard + Telegram"],
            ["Encargado barra", "Marcar barril terminado, registrar limpieza de líneas", "Solo Telegram (sin dashboard)"],
            ["Encargado cocina", "Lista de compras, reportar faltantes de insumos", "Solo Telegram (sin dashboard)"],
        ],
        col_widths=[3.5, 8.5, 4.5]
    )
    add_gap(doc, 80)

    # Funciones futuras
    add_h2(doc, "Funciones futuras previstas")
    add_data_table(doc,
        ["Tipo", "Función", "Notas"],
        [
            ["Operacional", "Marcar barril terminado (Telegram)", "Conecta directo con el menú: el tap se marca como agotado automáticamente."],
            ["Operacional", "Calendario de limpieza de líneas", "Regla configurable (ej: 14 días). Cron diario notifica si una línea no se ha limpiado."],
            ["Planificación", "Calendario de eventos y festividades cerveceras", "Cron semanal investiga festividades relevantes (LLM + web search). Estado 'sugerida' → apruebas/editas/rechazas por Telegram."],
            ["Planificación", "Calendario de ausencias de personal", "Vacaciones, días libres. Notificaciones para planear cuberturas."],
            ["Administrativa", "Lista de compras", "Encargado de cocina agrega insumos por Telegram (texto). Vista consolidada en dashboard."],
            ["Administrativa", "Gastos por Telegram (texto/voz/foto)", "Socio reporta gastos. Voz → transcripción (Whisper). Foto de ticket → LLM extrae monto/concepto. Categorización configurable."],
            ["Administrativa", "Reportes de gastos", "Tabla + totales por mes/categoría. Exportable."],
        ],
        col_widths=[3.0, 6.0, 7.5]
    )
    add_gap(doc, 120)

    # Fácil vs. con miga
    add_h2(doc, "Lo sencillo vs. lo que tiene miga")
    add_h3(doc, "Sencillo (mismo esfuerzo que Hermes)")
    add_bullet(doc, "Marcar barril terminado por Telegram")
    add_bullet(doc, "Lista de compras por texto")
    add_bullet(doc, "Gasto por texto con categoría")
    add_bullet(doc, "Reportes básicos (tabla + gráfico de totales por mes/categoría)")
    add_gap(doc, 80)
    add_h3(doc, "Con miga (no imposible, pero requiere iteración de calidad)")
    add_bullet(doc, [
        {"text": "Voz → texto: ", "bold": True},
        "transcripción (Whisper o similar). Funciona bien pero a veces equivoca nombres de cervezas o montos. Requiere confirmación antes de guardar.",
    ])
    add_bullet(doc, [
        {"text": "Foto de ticket → gasto estructurado: ", "bold": True},
        "visión LLM lee el ticket y extrae monto, fecha, concepto. Funciona en el 80-90% de los casos; los tickets oscuros/arrugados fallan. El flujo debe incluir revisión/edición antes de confirmar.",
    ])
    add_bullet(doc, [
        {"text": "Categorización automática: ", "bold": True},
        "hay que definir un catálogo de categorías contables con ustedes (no adivinarlo). Es un mini trabajo de setup, no técnico.",
    ])
    add_body(doc, "Nada de esto es bloqueante. Solo que las funciones de voz/foto no son 'magia que siempre sale perfecta' — son 'magia que a veces requiere confirmación', y eso se diseña así desde el inicio.")
    add_gap(doc, 120)

    # Orden de prioridades
    add_h2(doc, "Orden de prioridades interno sugerido")
    add_body(doc, "El dashboard crece en alcance. Cada sub-fase entrega valor de uso inmediato. No se construye todo de golpe.")
    add_bullet(doc, [{"text": "4a. ", "bold": True}, "Multi-usuario con roles (fundación; todo depende de esto)"])
    add_bullet(doc, [{"text": "4b. ", "bold": True}, "Marcar barril terminado (conecta directo con el menú, alto valor, simple)"])
    add_bullet(doc, [{"text": "4c. ", "bold": True}, "Calendario de eventos + festividades (conecta con sitio público)"])
    add_bullet(doc, [{"text": "4d. ", "bold": True}, "Limpieza de líneas con notificaciones"])
    add_bullet(doc, [{"text": "4e. ", "bold": True}, "Lista de compras"])
    add_bullet(doc, [{"text": "4f. ", "bold": True}, "Gastos por Telegram (texto primero, voz/foto después)"])
    add_bullet(doc, [{"text": "4g. ", "bold": True}, "Reportes de gastos"])


def build_section_16(doc):
    """Sección 16: Estrategia de producto - instancias personalizables"""
    add_h1(doc, "16 · Estrategia de producto: instancias personalizables", page_break_before=True)
    add_body(doc, [
        {"text": "Decisión: ", "bold": True},
        "producto base + instancias independientes por cliente (",
        {"text": "no multi-tenant", "bold": True, "color": P["accent"]},
        "). La intención a futuro es ofrecer sistemas similares a otros negocios (bares, cafeterías, restaurantes) partiendo de la base de Vecindario, adaptando módulos comunes y agregando particularidades.",
    ])
    add_body(doc, "Vecindario es la primera instancia y, al mismo tiempo, la fundación del producto base.")

    # Dos repos
    add_h2(doc, "Organización del código: 2 repos")
    add_body(doc, [
        "Desde el inicio se trabaja con dos repositorios Git separados:",
    ])
    add_bullet(doc, [
        {"text": "producto-base: ", "bold": True, "color": P["accent"]},
        "módulos comunes reusables (panel admin, auth+roles, dashboard, bot Telegram, API REST, módulos comunes como gastos/calendario/compras). Configurable pero no específico a un negocio.",
    ])
    add_bullet(doc, [
        {"text": "vecindario: ", "bold": True, "color": P["accent"]},
        "instancia que usa el producto base + sus particularidades (módulos específicos de bar de cerveza: limpieza líneas, rotación taps, festividades cerveceras). Configuración propia (nombre, logo, colores, dominio).",
    ])
    add_body(doc, "Cuando llegue un cliente nuevo: se clona producto-base, se configura, se seleccionan módulos, se agregan particularidades, se despliega como proyecto separado.")
    add_gap(doc, 120)

    # Comparación con multi-tenant
    add_h2(doc, "Diferencia con multi-tenant (lo que NO estamos haciendo)")
    add_data_table(doc,
        ["", "Multi-tenant", "Producto base + instancias (este proyecto)"],
        [
            ["Bases de datos", "1 compartida", "1 por cliente"],
            ["Logins", "Compartidos entre negocios", "Independientes por negocio"],
            ["Actualización de módulos comunes", "Afecta a todos a la vez", "Hay que mergear el cambio a cada instancia"],
            ["Personalización", "Limitada", "Total"],
            ["Esfuerzo inicial", "Alto", "Bajo"],
            ["Mantenimiento N clientes", "Centralizado", "Distribuido (N instancias)"],
        ],
        col_widths=[4.5, 6.0, 6.0]
    )
    add_gap(doc, 120)

    # Reglas de organización
    add_h2(doc, "Reglas de organización del código desde el inicio")
    add_body(doc, "Estas reglas son básicamente buenas prácticas de desarrollo; cuestan prácticamente lo mismo que desarrollar sin pensar en reutilizar, pero te ahorran dolores de cabeza cuando clones.")
    add_bullet(doc, [
        {"text": "Configuración en archivos, no hardcodeada: ", "bold": True},
        "el nombre 'Vecindario', los colores, el logo, el WhatsApp, los horarios, etc. viven en un archivo de config. Cuando clonas para la cafetería, cambias ese archivo y casi todo se adapta.",
    ])
    add_bullet(doc, [
        {"text": "/modules con cada función separada: ", "bold": True},
        "el código del dashboard organizado en carpetas por función (/modules/barriles, /modules/gastos, /modules/calendario). Para la cafetería, borras /modules/barriles y no afecta al resto.",
    ])
    add_bullet(doc, [
        {"text": "Nada de suposiciones del negocio embebidas: ", "bold": True},
        "por ejemplo, no asumir 'siempre hay 10 taps'. Ese '10' debería ser configurable, porque la cafetería no tiene taps, y otro bar puede tener 6.",
    ])
    add_bullet(doc, [
        {"text": "Cada módulo marcado como 'común' o 'específico de bar de cerveza': ", "bold": True},
        "cuando desarrollas una función para Vecindario, te preguntas: ¿esto es común a cualquier negocio similar, o es específico de un bar de cerveza? Los comunes van al producto base limpios y documentados.",
    ])
    add_gap(doc, 120)

    # El coste del modelo
    add_h2(doc, "El coste del modelo (y cuándo reconsiderar)")
    add_body(doc, [
        {"text": "El 'costo' de este modelo: ", "bold": True},
        "cuando arreglas un bug en un módulo común (ej: el módulo de gastos), tienes que mergear ese cambio a cada instancia. Con ",
        {"text": "pocos clientes (3-5) es manejable", "bold": True},
        ". Con ",
        {"text": "muchos (50+) se vuelve un problema", "bold": True},
        " — pero para cuando tengas 50 clientes, ya sabrás bien qué modelo te conviene y puedes migrar a multi-tenant si vale la pena.",
    ])
    add_body(doc, [
        {"text": "Conclusión: ", "bold": True},
        "no se construye multi-tenancy ahora. Se construye Vecindario bien, con código modular y configurable, y se evalúa la migración a multi-tenant solo cuando el volumen de clientes lo justifique.",
    ])


def move_elements_before_closing(doc):
    """
    python-docx appends new elements at the end of the body, AFTER the
    closing paragraph ("Tómate tu tiempo..."). We need to move our newly
    added elements to BEFORE the closing paragraph.

    The closing paragraph is the second-to-last element (last is sectPr).
    """
    body = doc.element.body
    children = list(body)
    # The sectPr is the last child
    sectPr = children[-1]
    # The closing paragraph is the one containing "Tómate tu tiempo" - it should be
    # the second-to-last NOW, but we just added many elements after it.
    # Find the closing paragraph by text.
    closing_p = None
    for ch in children:
        if ch.tag == qn("w:p"):
            texts = ch.findall(".//" + qn("w:t"))
            txt = "".join(t.text or "" for t in texts)
            if "Tómate tu tiempo" in txt:
                closing_p = ch
                break
    if closing_p is None:
        print("WARNING: closing paragraph not found, leaving elements at end")
        return
    # All paragraphs/tables added by python-docx appear after the original
    # closing paragraph (which was the last paragraph before sectPr).
    # We need to move them to BEFORE the closing paragraph.
    # Strategy: find closing_p index, move all elements after it (except sectPr)
    # to be inserted before it.

    # Identify elements to move: everything between closing_p and sectPr
    closing_idx = children.index(closing_p)
    elements_to_move = []
    for ch in children[closing_idx + 1:]:
        if ch.tag == qn("w:sectPr"):
            continue  # don't move sectPr
        elements_to_move.append(ch)

    # Remove them from current position
    for ch in elements_to_move:
        body.remove(ch)

    # Re-find closing_p (its position may have shifted after removals, but
    # actually it shouldn't because we removed elements AFTER it)
    # Insert elements before closing_p
    for ch in elements_to_move:
        closing_p.addprevious(ch)

    print(f"Moved {len(elements_to_move)} elements before closing paragraph")


def main():
    doc = Document(DOC_PATH)
    print(f"Opened: {DOC_PATH}")
    print(f"  Paragraphs before: {len(doc.paragraphs)}")
    print(f"  Tables before: {len(doc.tables)}")

    # Build sections (appends at end, before sectPr)
    build_section_15(doc)
    build_section_16(doc)

    print(f"  Paragraphs after build: {len(doc.paragraphs)}")
    print(f"  Tables after build: {len(doc.tables)}")

    # Move new elements to before the closing paragraph
    move_elements_before_closing(doc)

    doc.save(DOC_PATH)
    print(f"Saved: {DOC_PATH}")
    print(f"  Final paragraphs: {len(doc.paragraphs)}")
    print(f"  Final tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
